#!/usr/bin/env python3
"""ATS / Workday-optimised Word (.docx) versions of Gabe's two CV variants.
Structure tuned so resume parsers (Workday, Greenhouse) capture every field:
- Each role: Title / Company / Location / Dates on their own lines, then the description as
  plain hyphen bullets directly underneath (Word list styles get dropped by Workday).
- Skills, clients and tools sit AFTER all experience so they do not bleed into a role.
Single column, standard font, no tables/text boxes/columns. Content from cv_data.py. No em dashes."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
import cv_data as D

INK=RGBColor(0x1a,0x1a,0x1a); TEAL=RGBColor(0x0c,0x4f,0x4d)
def clean(t):
    t=re.sub(r'\*\*(.+?)\*\*',r'\1',t)
    return re.sub(r'<[^>]+>','',t).replace('&amp;','&').replace('&nbsp;',' ').replace('•','-')
def para(doc,txt,size=10.5,bold=False,color=INK,after=2,before=0):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after); p.paragraph_format.space_before=Pt(before)
    r=p.add_run(clean(txt)); r.font.size=Pt(size); r.font.bold=bold; r.font.color.rgb=color; r.font.name="Calibri"
    return p
def heading(doc,txt):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(txt.upper()); r.font.size=Pt(11); r.font.bold=True; r.font.color.rgb=TEAL; r.font.name="Calibri"

def build(key,v):
    doc=Document()
    for s in doc.sections:
        s.top_margin=s.bottom_margin=Inches(0.5); s.left_margin=s.right_margin=Inches(0.7)
    # header
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(0)
    r=p.add_run(D.NAME); r.font.size=Pt(20); r.font.bold=True; r.font.color.rgb=TEAL; r.font.name="Calibri"
    para(doc,v["subtitle"],size=10.5,bold=True,color=INK,after=2)
    para(doc,D.CONTACT,size=9.5,after=0)
    para(doc,D.RTW,size=9.5,after=2)
    # summary
    heading(doc,"Summary"); para(doc,v["profile"],size=10,after=2)
    # experience (parser-friendly: each field on its own line, plain hyphen bullets directly under)
    heading(doc,"Work Experience")
    for title,co,date,bl in D.EXPERIENCE:
        comp,loc=(co.rsplit(", ",1)+[""])[:2] if ", " in co else (co,"")
        para(doc,title,size=11,bold=True,color=INK,after=0,before=6)
        para(doc,comp,size=10,bold=True,color=INK,after=0)
        if loc: para(doc,loc,size=10,after=0)
        para(doc,date,size=10,after=2)
        for b in bl: para(doc,"- "+clean(b),size=10,after=1)
    # everything else AFTER experience so it cannot be grabbed as a role description
    heading(doc,"Key Achievements")
    for h in v.get("highlights",D.ACHIEVEMENTS): para(doc,"- "+clean(h),size=10,after=1)
    heading(doc,"Skills"); para(doc,", ".join(v["skills"]),size=10,after=2)
    heading(doc,"Selected Clients"); para(doc,D.BRANDS,size=10,after=2)
    heading(doc,"Education"); para(doc,D.EDUCATION,size=10,after=1)
    heading(doc,"Languages"); para(doc,D.LANGUAGES,size=10,after=1)
    heading(doc,"Tools"); para(doc,D.TOOLS,size=10,after=1)
    out=v["file"]+".docx"; doc.save(out); print("wrote",out)

if __name__=="__main__":
    for k,v in D.VARIANTS.items(): build(k,v)
